import io
import json
import os

import fitz
import pytesseract
from docx import Document
from flask import Flask, Response, jsonify, request
from openpyxl import load_workbook
from PIL import Image, ImageDraw
import xlrd
from xlutils.copy import copy as copy_xls

app = Flask(__name__)
MAX_BYTES = int(os.getenv("PRESIDIO_MAX_FILE_BYTES", str(50 << 20)))


def uploaded():
    item = request.files.get("file")
    if item is None:
        raise ValueError("file is required")
    data = item.read(MAX_BYTES + 1)
    if len(data) > MAX_BYTES:
        raise ValueError("file is too large")
    return data, request.form.get("media_type", "application/octet-stream").lower()


def kind(media_type):
    if media_type in ("text/plain", "text/csv", "application/csv"): return "text"
    if media_type == "application/pdf": return "pdf"
    if media_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",): return "docx"
    if media_type in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",): return "xlsx"
    if media_type in ("application/vnd.ms-excel", "application/xls"): return "xls"
    if media_type.startswith("image/"): return "image"
    raise ValueError(f"unsupported media type: {media_type}")


def image_words(image):
    data = pytesseract.image_to_data(image, lang=os.getenv("PRESIDIO_OCR_LANG", "eng+fra"), output_type=pytesseract.Output.DICT)
    words = []
    for index, text in enumerate(data["text"]):
        text = text.strip()
        if text:
            words.append((f"word:{index}", text, (data["left"][index], data["top"][index], data["width"][index], data["height"][index])))
    return words


def extract(data, media_type):
    file_kind = kind(media_type)
    if file_kind == "text":
        text = data.decode("utf-8-sig")
        return [{"id": "text:0", "text": text}] if text else []
    if file_kind == "docx":
        doc = Document(io.BytesIO(data)); result = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text: result.append({"id": f"p:{i}", "text": paragraph.text})
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if cell.text: result.append({"id": f"c:{ti}:{ri}:{ci}", "text": cell.text})
        return result
    if file_kind == "xlsx":
        book = load_workbook(io.BytesIO(data)); result = []
        for sheet in book.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value:
                        result.append({"id": f"cell:{sheet.title}:{cell.coordinate}", "text": cell.value})
        return result
    if file_kind == "xls":
        book = xlrd.open_workbook(file_contents=data, formatting_info=True); result = []
        for si, sheet in enumerate(book.sheets()):
            for row in range(sheet.nrows):
                for col in range(sheet.ncols):
                    value = sheet.cell_value(row, col)
                    if isinstance(value, str) and value:
                        result.append({"id": f"xls:{si}:{row}:{col}", "text": value})
        return result
    if file_kind == "pdf":
        doc = fitz.open(stream=data, filetype="pdf"); result = []
        for page_index, page in enumerate(doc):
            for block_index, block in enumerate(page.get_text("blocks")):
                text = block[4].strip()
                if text: result.append({"id": f"block:{page_index}:{block_index}", "text": text})
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            for ident, text, _ in image_words(image): result.append({"id": f"ocr:{page_index}:{ident}", "text": text})
        return result
    image = Image.open(io.BytesIO(data)).convert("RGB")
    return [{"id": ident, "text": text} for ident, text, _ in image_words(image)]


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]: run.text = ""
    else: paragraph.add_run(text)


def render(data, media_type, replacements):
    values = {item["id"]: item["text"] for item in replacements}
    file_kind = kind(media_type)
    if file_kind == "text":
        return values.get("text:0", "").encode("utf-8")
    if file_kind == "docx":
        doc = Document(io.BytesIO(data))
        for i, paragraph in enumerate(doc.paragraphs):
            if f"p:{i}" in values: set_paragraph_text(paragraph, values[f"p:{i}"])
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    ident = f"c:{ti}:{ri}:{ci}"
                    if ident in values: cell.text = values[ident]
        out = io.BytesIO(); doc.save(out); return out.getvalue()
    if file_kind == "xlsx":
        book = load_workbook(io.BytesIO(data))
        for sheet in book.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    ident = f"cell:{sheet.title}:{cell.coordinate}"
                    if ident in values: cell.value = values[ident]
        out = io.BytesIO(); book.save(out); return out.getvalue()
    if file_kind == "xls":
        source = xlrd.open_workbook(file_contents=data, formatting_info=True)
        target = copy_xls(source)
        for si, sheet in enumerate(source.sheets()):
            writable = target.get_sheet(si)
            for row in range(sheet.nrows):
                for col in range(sheet.ncols):
                    ident = f"xls:{si}:{row}:{col}"
                    if ident in values: writable.write(row, col, values[ident])
        out = io.BytesIO(); target.save(out); return out.getvalue()
    if file_kind == "pdf":
        doc = fitz.open(stream=data, filetype="pdf")
        for page_index, page in enumerate(doc):
            blocks = page.get_text("blocks")
            for block_index, block in enumerate(blocks):
                ident = f"block:{page_index}:{block_index}"
                if ident in values and values[ident] != block[4].strip():
                    rect = fitz.Rect(block[:4]); page.add_redact_annot(rect, fill=(1, 1, 1), text=values[ident], fontsize=8)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            for ident, text, (left, top, width, height) in image_words(image):
                full_id = f"ocr:{page_index}:{ident}"
                if full_id in values and values[full_id] != text:
                    page.add_redact_annot(fitz.Rect(left / 1.5, top / 1.5, (left + width) / 1.5, (top + height) / 1.5), fill=(0, 0, 0))
            page.apply_redactions()
        return doc.tobytes(garbage=4, deflate=True)
    source_image = Image.open(io.BytesIO(data)); output_format = source_image.format or "PNG"
    image = source_image.convert("RGB"); draw = ImageDraw.Draw(image)
    for ident, text, (left, top, width, height) in image_words(image):
        if ident in values and values[ident] != text: draw.rectangle((left, top, left + width, top + height), fill="black")
    out = io.BytesIO(); image.save(out, format=output_format); return out.getvalue()


@app.get("/readyz")
def ready(): return jsonify({"status": "ready"})


@app.post("/v1/extract")
def extract_route():
    try:
        data, media_type = uploaded()
        return jsonify({"segments": extract(data, media_type)})
    except Exception as exc:
        app.logger.exception("Presidio extract failed for media_type=%s size=%d", request.form.get("media_type", "unknown"), request.content_length or 0)
        return jsonify({"error": str(exc)}), 422


@app.post("/v1/render")
def render_route():
    try:
        data, media_type = uploaded()
        replacements = json.loads(request.form.get("replacements", "[]"))
        return Response(render(data, media_type, replacements), mimetype=media_type)
    except Exception as exc:
        app.logger.exception("Presidio render failed for media_type=%s size=%d", request.form.get("media_type", "unknown"), request.content_length or 0)
        return jsonify({"error": str(exc)}), 422
